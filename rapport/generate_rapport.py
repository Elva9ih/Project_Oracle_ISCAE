"""Generate the project rapport as a Word document."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

SCREENSHOTS = os.path.join(os.path.dirname(__file__), "screenshots")
OUTPUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Rapport_Projet_Oracle.docx")

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(26, 122, 58)
    h.font.name = 'Calibri'

def add_screenshot(name, caption, width=5.8):
    path = os.path.join(SCREENSHOTS, f"{name}.png")
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

# ============================================================
# PAGE DE GARDE
# ============================================================
doc.add_paragraph()

# Logo ISCAE
logo_path = os.path.join(os.path.dirname(__file__), "logo-iscae.png")
if os.path.exists(logo_path):
    doc.add_picture(logo_path, width=Inches(2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ISCAE - Nouakchott")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(26, 122, 58)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Master MPIAG - 2025/2026")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROJET ORACLE")
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = RGBColor(26, 122, 58)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Oracle dans un Écosystème Applicatif")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gestion des Absences Étudiants")
r.font.size = Pt(20)
r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Application Web Django + Base de données Oracle 18c XE")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Réalisé par :")
r.font.size = Pt(12)
r.font.bold = True
r.font.color.rgb = RGBColor(26, 122, 58)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("IP22240 — Cherifa Mohamed Yeslem Bellahi\nIP22237 — El Vagih Ahmed Zeine")
r.font.size = Pt(13)
r.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Module : Oracle dans un Écosystème Applicatif")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============================================================
# TABLE DES MATIÈRES
# ============================================================
doc.add_heading("Table des matières", level=1)
toc_items = [
    "1. Introduction",
    "2. Architecture du projet",
    "3. Base de données Oracle",
    "   3.1. Modèle conceptuel (MCD)",
    "   3.2. Schéma relationnel (Tables et Contraintes)",
    "   3.3. Index et Optimisation",
    "   3.4. PL/SQL : Packages, Fonctions et Procédures",
    "   3.5. Triggers (Audit et Règles métier)",
    "   3.6. Vues et Vues Matérialisées",
    "   3.7. Transactions et Gestion de la concurrence",
    "   3.8. Génération de données volumineuses",
    "   3.9. Plans d'exécution (EXPLAIN PLAN)",
    "   3.10. Export et Sauvegarde (Data Pump)",
    "4. Application Django",
    "   4.1. Architecture technique",
    "   4.2. Connexion Django-Oracle",
    "   4.3. Interfaces par rôle",
    "5. Captures d'écran",
    "6. Conclusion",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "Ce rapport présente le projet réalisé dans le cadre du module « Oracle dans un Écosystème Applicatif » "
    "du Master MPIAG à l'ISCAE. Le projet consiste à concevoir et développer une application complète de "
    "gestion des absences étudiants, mettant en œuvre les fonctionnalités avancées d'Oracle Database 18c XE "
    "intégrées dans une application web Django."
)
doc.add_paragraph(
    "L'objectif est de démontrer la maîtrise des concepts fondamentaux et avancés d'Oracle : "
    "modélisation relationnelle avec contraintes d'intégrité, programmation PL/SQL (packages, fonctions, "
    "procédures), triggers d'audit et métier, vues et vues matérialisées, optimisation des requêtes "
    "(EXPLAIN PLAN, index), gestion des transactions, et export/sauvegarde via Data Pump."
)
doc.add_paragraph(
    "Le domaine choisi — la gestion des absences — est pertinent car il implique naturellement des règles "
    "métier complexes (seuils d'exclusion, justifications, notifications automatiques) qui se prêtent "
    "parfaitement à l'utilisation de PL/SQL et des triggers Oracle."
)

doc.add_heading("Technologies utilisées", level=2)
add_table(
    ["Composant", "Technologie", "Version"],
    [
        ["Base de données", "Oracle Database XE", "18c (18.4.0)"],
        ["Langage BD", "PL/SQL", "-"],
        ["Framework Web", "Django", "3.2.25 LTS"],
        ["Langage Backend", "Python", "3.11.2"],
        ["Driver Oracle", "python-oracledb", "3.4.2"],
        ["Frontend", "Bootstrap 5 + Chart.js", "5.3.2 / 4.4.1"],
        ["IDE Base de données", "SQL Developer", "-"],
        ["OS", "Windows 11", "-"],
    ]
)

doc.add_page_break()

# ============================================================
# 2. ARCHITECTURE
# ============================================================
doc.add_heading("2. Architecture du projet", level=1)
doc.add_paragraph(
    "Le projet suit une architecture 3-tiers classique avec une répartition 75% Oracle / 25% Django, "
    "conformément aux exigences du cahier des charges :"
)
doc.add_paragraph("• Couche Données (Oracle 18c) : Stockage, contraintes, logique métier PL/SQL, triggers, vues")
doc.add_paragraph("• Couche Application (Django) : Interface web, appels aux procédures Oracle, gestion des utilisateurs")
doc.add_paragraph("• Couche Présentation (Bootstrap 5) : Interface responsive, graphiques Chart.js, exports PDF/Excel")

doc.add_heading("Principes architecturaux", level=2)
doc.add_paragraph(
    "La logique métier est volontairement placée côté Oracle (packages PL/SQL) et non dans Django. "
    "L'application Django agit comme un client qui appelle les procédures et fonctions Oracle "
    "via le module oracle_services.py. Cette approche garantit l'intégrité des données indépendamment "
    "du client utilisé et démontre la puissance de PL/SQL."
)

doc.add_heading("Structure des fichiers", level=2)
doc.add_paragraph(
    "oracle_scripts/          — 7 scripts SQL exécutés en séquence\n"
    "  01_schema.sql           — Tables, contraintes, index\n"
    "  02_plsql.sql            — Packages, fonctions, procédures\n"
    "  03_triggers.sql         — Triggers d'audit et métier\n"
    "  04_views.sql            — Vues et vues matérialisées\n"
    "  05_data_generation.sql  — Données volumineuses (120K+ lignes)\n"
    "  06_optimization.sql     — EXPLAIN PLAN et benchmarks\n"
    "  07_backup_export.sql    — Configuration Data Pump\n"
    "django_app/              — Application web complète\n"
    "  absences/models.py      — Modèles Django (managed=False)\n"
    "  absences/oracle_services.py — Appels PL/SQL\n"
    "  absences/views.py       — Vues avec contrôle de rôles\n"
    "  templates/              — 15+ templates HTML",
    style='No Spacing'
)

doc.add_page_break()

# ============================================================
# 3. BASE DE DONNÉES ORACLE
# ============================================================
doc.add_heading("3. Base de données Oracle", level=1)

# 3.1 MCD
doc.add_heading("3.1. Modèle Conceptuel de Données (MCD)", level=2)
doc.add_paragraph(
    "Le système de gestion des absences repose sur les entités suivantes et leurs relations :"
)
add_table(
    ["Entité", "Description", "Cardinalité principale"],
    [
        ["FILIERES", "Filières d'études (MPIAG, GFC, MRH...)", "1 filière → N groupes"],
        ["GROUPES", "Groupes d'étudiants par filière", "1 groupe → N étudiants"],
        ["SEMESTRES", "Périodes académiques (S1, S2)", "1 semestre → N matières"],
        ["ETUDIANTS", "Dossiers étudiants avec statut", "1 étudiant → N absences"],
        ["ENSEIGNANTS", "Corps enseignant", "1 enseignant → N séances"],
        ["MATIERES", "Matières enseignées par filière/semestre", "1 matière → N séances"],
        ["AFFECTATIONS", "Liaison enseignant-matière-groupe", "Table d'association"],
        ["SEANCES", "Séances de cours planifiées/effectuées", "1 séance → N absences"],
        ["ABSENCES", "Enregistrement des absences", "1 absence → 0..N justificatifs"],
        ["JUSTIFICATIFS", "Documents justificatifs soumis", "Lié à une absence"],
        ["SEUILS_ABSENCES", "Configuration des seuils par filière", "1 filière → 1 seuil/an"],
        ["NOTIFICATIONS", "Alertes automatiques aux étudiants", "Générées par triggers"],
        ["AUDIT_LOG", "Journal d'audit automatique", "Alimenté par triggers"],
    ]
)

# 3.2 Schema
doc.add_heading("3.2. Schéma Relationnel (Tables et Contraintes)", level=2)
doc.add_paragraph(
    "Le schéma comprend 13 tables avec un ensemble complet de contraintes d'intégrité :"
)

doc.add_heading("Contraintes implémentées", level=3)
add_table(
    ["Type de contrainte", "Nombre", "Exemples"],
    [
        ["PRIMARY KEY", "13", "Chaque table a une PK auto-générée (IDENTITY)"],
        ["FOREIGN KEY", "15", "fk_etudiant_groupe, fk_absence_seance, etc."],
        ["UNIQUE", "10", "cne, matricule, code_filiere, (id_etudiant, id_seance)"],
        ["CHECK", "12", "statut IN ('ACTIF','EXCLU'...), coefficient > 0, est_justifiee IN (0,1)"],
        ["NOT NULL", "40+", "Sur tous les champs obligatoires"],
        ["DEFAULT", "8+", "SYSDATE, 'ACTIF', 0, 'PLANIFIEE'"],
    ]
)

doc.add_paragraph()
doc.add_heading("Exemple : Table ABSENCES", level=3)
p = doc.add_paragraph()
r = p.add_run(
    "CREATE TABLE ABSENCES (\n"
    "    id_absence    NUMBER GENERATED ALWAYS AS IDENTITY PRIMARY KEY,\n"
    "    id_etudiant   NUMBER        NOT NULL,\n"
    "    id_seance     NUMBER        NOT NULL,\n"
    "    date_absence  DATE          DEFAULT SYSDATE NOT NULL,\n"
    "    est_justifiee NUMBER(1)     DEFAULT 0 NOT NULL,\n"
    "    motif         VARCHAR2(200),\n"
    "    saisi_par     VARCHAR2(50),\n"
    "    CONSTRAINT fk_absence_etudiant FOREIGN KEY (id_etudiant)\n"
    "        REFERENCES ETUDIANTS(id_etudiant),\n"
    "    CONSTRAINT fk_absence_seance FOREIGN KEY (id_seance)\n"
    "        REFERENCES SEANCES(id_seance),\n"
    "    CONSTRAINT uk_absence UNIQUE (id_etudiant, id_seance),\n"
    "    CONSTRAINT chk_est_justifiee CHECK (est_justifiee IN (0, 1))\n"
    ");"
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

# 3.3 Index
doc.add_heading("3.3. Index et Optimisation", level=2)
doc.add_paragraph("22 index ont été créés pour optimiser les performances des requêtes :")
add_table(
    ["Type d'index", "Nombre", "Exemples"],
    [
        ["Index sur FK", "14", "idx_absence_etudiant, idx_seance_matiere"],
        ["Index de recherche", "4", "idx_seance_date, idx_absence_date"],
        ["Index fonctionnel", "2", "idx_etudiant_nom (UPPER(nom), UPPER(prenom))"],
        ["Index de filtrage", "2", "idx_etudiant_statut, idx_absence_justifiee"],
    ]
)

# 3.4 PL/SQL
doc.add_page_break()
doc.add_heading("3.4. PL/SQL : Packages, Fonctions et Procédures", level=2)
doc.add_paragraph(
    "Toute la logique métier est encapsulée dans deux packages PL/SQL, conformément aux bonnes pratiques Oracle :"
)

doc.add_heading("Package PKG_ABSENCES", level=3)
doc.add_paragraph("Ce package gère l'ensemble du cycle de vie des absences :")
add_table(
    ["Élément", "Type", "Description"],
    [
        ["GET_NB_ABSENCES", "Fonction", "Retourne le nombre total d'absences d'un étudiant (optionnel: par semestre)"],
        ["GET_NB_ABSENCES_NON_JUSTIFIEES", "Fonction", "Nombre d'absences non justifiées"],
        ["GET_TAUX_ABSENCE", "Fonction", "Calcule le taux d'absence en % (absences/séances effectuées)"],
        ["EST_EN_DEPASSEMENT", "Fonction", "Vérifie si l'étudiant dépasse les seuils (NORMAL/AVERTISSEMENT/EXCLUSION)"],
        ["GET_STATUT_ETUDIANT_ABSENCES", "Fonction", "Retourne un message détaillé avec le taux et le niveau d'alerte"],
        ["MARQUER_ABSENCE", "Procédure", "Enregistre une absence avec vérifications (étudiant actif, pas de doublon, séance existante)"],
        ["MARQUER_ABSENCES_BULK", "Procédure", "Marquage en masse via CSV d'IDs avec SAVEPOINT/ROLLBACK"],
        ["JUSTIFIER_ABSENCE", "Procédure", "Soumet un justificatif pour une absence"],
        ["TRAITER_JUSTIFICATIF", "Procédure", "Accepte/refuse un justificatif et met à jour l'absence automatiquement"],
        ["VERIFIER_SEUILS_ETUDIANT", "Procédure", "Vérifie les seuils et génère des notifications/exclusions automatiques"],
        ["CLOTURER_SEMESTRE", "Procédure", "Clôture un semestre (séances → effectuées, justificatifs → refusés)"],
    ]
)

doc.add_paragraph()
doc.add_heading("Exemple : Fonction GET_TAUX_ABSENCE", level=3)
p = doc.add_paragraph()
r = p.add_run(
    "FUNCTION get_taux_absence(p_id_etudiant NUMBER,\n"
    "    p_id_semestre NUMBER DEFAULT NULL) RETURN NUMBER IS\n"
    "    v_total NUMBER; v_abs NUMBER;\n"
    "BEGIN\n"
    "    SELECT COUNT(*) INTO v_total\n"
    "    FROM SEANCES s JOIN ETUDIANTS e ON s.id_groupe = e.id_groupe\n"
    "    WHERE e.id_etudiant = p_id_etudiant\n"
    "      AND s.statut = 'EFFECTUEE';\n"
    "    IF v_total = 0 THEN RETURN 0; END IF;\n"
    "    v_abs := get_nb_absences(p_id_etudiant, p_id_semestre);\n"
    "    RETURN ROUND((v_abs / v_total) * 100, 2);\n"
    "END;"
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

doc.add_paragraph()
doc.add_heading("Package PKG_STATISTIQUES", level=3)
doc.add_paragraph("Fournit des curseurs de données pour les rapports et tableaux de bord :")
doc.add_paragraph("• GET_TOP_ABSENTS : Top N étudiants les plus absents (avec SYS_REFCURSOR)")
doc.add_paragraph("• GET_STATS_PAR_MATIERE : Statistiques d'absences par matière")
doc.add_paragraph("• GET_STATS_PAR_GROUPE : Statistiques par groupe avec moyenne par étudiant")
doc.add_paragraph("• GET_EVOLUTION_MENSUELLE : Évolution mensuelle des absences")
doc.add_paragraph("• GENERER_RAPPORT_ABSENCES : Rapport détaillé par groupe et période")

# 3.5 Triggers
doc.add_page_break()
doc.add_heading("3.5. Triggers (Audit et Règles métier)", level=2)
doc.add_paragraph("8 triggers assurent l'audit automatique et l'application des règles métier :")

doc.add_heading("Triggers d'audit", level=3)
add_table(
    ["Trigger", "Table", "Description"],
    [
        ["TRG_AUDIT_ABSENCES", "ABSENCES", "Journalise chaque INSERT/UPDATE/DELETE avec anciennes et nouvelles valeurs dans AUDIT_LOG"],
        ["TRG_AUDIT_ETUDIANTS", "ETUDIANTS", "Trace toute modification de dossier étudiant"],
        ["TRG_AUDIT_JUSTIFICATIFS", "JUSTIFICATIFS", "Trace la soumission et le traitement des justificatifs"],
    ]
)

doc.add_paragraph()
doc.add_heading("Triggers métier", level=3)
add_table(
    ["Trigger", "Type", "Description"],
    [
        ["TRG_ABSENCE_DATE_AUTO", "BEFORE INSERT", "Copie automatiquement la date de la séance dans l'absence"],
        ["TRG_PROTECT_ABSENCE_JUSTIFIEE", "BEFORE UPDATE", "Empêche de retirer la justification d'une absence déjà justifiée"],
        ["TRG_CHECK_ETUDIANT_ACTIF", "BEFORE INSERT", "Bloque l'insertion d'absence pour un étudiant non actif"],
        ["TRG_NOTIF_CHANGEMENT_STATUT", "AFTER UPDATE", "Crée une notification automatique quand le statut d'un étudiant change"],
        ["TRG_SEANCE_EFFECTUEE", "AFTER INSERT", "Marque automatiquement une séance comme effectuée dès qu'une absence y est saisie"],
    ]
)

doc.add_paragraph()
doc.add_heading("Exemple : Trigger d'audit", level=3)
p = doc.add_paragraph()
r = p.add_run(
    "CREATE OR REPLACE TRIGGER trg_audit_absences\n"
    "AFTER INSERT OR UPDATE OR DELETE ON ABSENCES\n"
    "FOR EACH ROW\n"
    "BEGIN\n"
    "    IF INSERTING THEN\n"
    "        INSERT INTO AUDIT_LOG (table_name, operation,\n"
    "            record_id, new_values)\n"
    "        VALUES ('ABSENCES', 'INSERT', :NEW.id_absence,\n"
    "            'id_etudiant=' || :NEW.id_etudiant);\n"
    "    ELSIF UPDATING THEN\n"
    "        INSERT INTO AUDIT_LOG (table_name, operation,\n"
    "            record_id, old_values, new_values)\n"
    "        VALUES ('ABSENCES', 'UPDATE', :NEW.id_absence,\n"
    "            'est_justifiee=' || :OLD.est_justifiee,\n"
    "            'est_justifiee=' || :NEW.est_justifiee);\n"
    "    END IF;\n"
    "END;"
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

# 3.6 Views
doc.add_page_break()
doc.add_heading("3.6. Vues et Vues Matérialisées", level=2)

doc.add_heading("Vues classiques (5)", level=3)
add_table(
    ["Vue", "Description"],
    [
        ["V_ABSENCES_DETAIL", "Jointure complète : absence + étudiant + groupe + filière + séance + matière + enseignant"],
        ["V_RESUME_ABSENCES_ETUDIANT", "Résumé par étudiant : total, justifiées, non justifiées, taux"],
        ["V_SEANCES_ABSENTS", "Chaque séance avec le nombre d'absents et le taux d'absence"],
        ["V_JUSTIFICATIFS_EN_ATTENTE", "Justificatifs en attente de traitement avec info étudiant/matière"],
        ["V_NOTIFICATIONS_NON_LUES", "Notifications non lues avec info étudiant"],
    ]
)

doc.add_paragraph()
doc.add_heading("Vues Matérialisées (3 + 1 vue KPI)", level=3)
doc.add_paragraph(
    "Les vues matérialisées stockent physiquement les résultats des requêtes complexes. "
    "Elles sont rafraîchies à la demande (REFRESH COMPLETE ON DEMAND) via la procédure REFRESH_ALL_MV, "
    "appelable depuis l'interface Django."
)
add_table(
    ["Vue Matérialisée", "Contenu", "Usage"],
    [
        ["MV_STATS_PAR_GROUPE", "Nb étudiants, absences totales/justifiées/non justifiées, moyenne par étudiant", "Dashboard + page stats groupes"],
        ["MV_STATS_PAR_MATIERE", "Nb séances, absences, moyenne par séance pour chaque matière", "Page stats matières"],
        ["MV_EVOLUTION_MENSUELLE", "Total absences par mois, ventilé justifiées/non justifiées", "Graphique Chart.js du dashboard"],
        ["V_KPI_GLOBAL (vue)", "KPIs : étudiants actifs/exclus, enseignants, séances, absences, justificatifs en attente", "Cartes KPI du dashboard"],
    ]
)

# 3.7 Transactions
doc.add_heading("3.7. Transactions et Gestion de la concurrence", level=2)
doc.add_paragraph("Le projet utilise les mécanismes transactionnels Oracle suivants :")
doc.add_paragraph("• COMMIT : Validation explicite après chaque opération métier réussie")
doc.add_paragraph("• ROLLBACK : Annulation en cas d'erreur dans les procédures")
doc.add_paragraph("• SAVEPOINT : Points de reprise dans les opérations en masse (MARQUER_ABSENCES_BULK, CLOTURER_SEMESTRE)")
doc.add_paragraph("• Gestion d'exceptions : RAISE_APPLICATION_ERROR avec codes d'erreur personnalisés (-20001 à -20021)")
doc.add_paragraph(
    "Exemple : La procédure MARQUER_ABSENCES_BULK utilise un SAVEPOINT avant de traiter la liste CSV "
    "d'étudiants. Si aucune absence n'a pu être enregistrée, un ROLLBACK TO SAVEPOINT annule toutes "
    "les tentatives partielles, garantissant l'atomicité de l'opération."
)

# 3.8 Data generation
doc.add_heading("3.8. Génération de données volumineuses", level=2)
doc.add_paragraph(
    "Pour tester les performances et démontrer la capacité de traitement, un script PL/SQL génère "
    "un volume significatif de données :"
)
add_table(
    ["Table", "Nombre de lignes", "Méthode"],
    [
        ["FILIERES", "5", "INSERT direct"],
        ["GROUPES", "15", "Boucle PL/SQL (3 par filière)"],
        ["SEMESTRES", "2", "INSERT direct"],
        ["ENSEIGNANTS", "50", "Boucle avec tableaux PL/SQL de noms marocains"],
        ["MATIERES", "50", "10 par filière, réparties sur 2 semestres"],
        ["ETUDIANTS", "795", "53 par groupe × 15 groupes"],
        ["AFFECTATIONS", "150", "Enseignant-matière-groupe (round-robin)"],
        ["SEANCES", "3 000", "20 séances par affectation"],
        ["ABSENCES", "120 000", "Sélection aléatoire étudiant+séance avec DBMS_RANDOM"],
        ["JUSTIFICATIFS", "~24 000", "Pour les 20% d'absences justifiées"],
        ["TOTAL", "~148 000+", ""],
    ]
)
doc.add_paragraph(
    "Technique d'optimisation : Les triggers d'audit sont désactivés pendant l'insertion en masse "
    "(ALTER TRIGGER ... DISABLE) puis réactivés, réduisant le temps d'insertion de plusieurs minutes."
)

# 3.9 Explain Plan
doc.add_heading("3.9. Plans d'exécution (EXPLAIN PLAN)", level=2)
doc.add_paragraph(
    "Le script 06_optimization.sql exécute EXPLAIN PLAN sur les requêtes principales pour analyser "
    "et démontrer l'utilisation des index :"
)
doc.add_paragraph("• Recherche d'absences par CNE étudiant → utilisation de l'index unique sur CNE")
doc.add_paragraph("• Top 10 étudiants absents → index sur est_justifiee + GROUP BY optimisé")
doc.add_paragraph("• Statistiques par matière → jointures optimisées par index FK")
doc.add_paragraph("• Requête sur vue matérialisée → accès direct sans jointure (MAT_VIEW ACCESS FULL)")
doc.add_paragraph("• Comparaison requête directe vs vue matérialisée → démonstration du gain de performance")
doc.add_paragraph("• Recherche par date → utilisation de idx_absence_date (INDEX RANGE SCAN)")
doc.add_paragraph("• Recherche par nom → utilisation de l'index fonctionnel sur UPPER(nom)")

# 3.10 Backup
doc.add_heading("3.10. Export et Sauvegarde (Data Pump)", level=2)
doc.add_paragraph("La stratégie de sauvegarde est configurée dans le script 07_backup_export.sql :")
doc.add_paragraph("• Data Pump Export (expdp) : Export complet du schéma gestion_absences")
doc.add_paragraph("• Data Pump Import (impdp) : Restauration complète du schéma")
doc.add_paragraph("• Script batch Windows : Sauvegarde automatisée avec horodatage et rotation (30 jours)")
doc.add_paragraph("• Planification : Configuration via schtasks pour exécution quotidienne à 02h00")

doc.add_page_break()

# ============================================================
# 4. APPLICATION DJANGO
# ============================================================
doc.add_heading("4. Application Django", level=1)

doc.add_heading("4.1. Architecture technique", level=2)
doc.add_paragraph(
    "L'application Django utilise le pattern MVT (Model-View-Template) avec une couche de service "
    "(oracle_services.py) pour les appels PL/SQL :"
)
doc.add_paragraph("• Models (managed=False) : Mapping des tables Oracle existantes sans migration Django")
doc.add_paragraph("• Views : Logique de contrôle avec décorateurs de rôles (@admin_required, @enseignant_required)")
doc.add_paragraph("• Templates : Interface Bootstrap 5 responsive avec sidebar dynamique selon le rôle")
doc.add_paragraph("• oracle_services.py : Couche d'abstraction qui appelle cursor.callfunc() et cursor.callproc()")

doc.add_heading("4.2. Connexion Django-Oracle", level=2)
p = doc.add_paragraph()
r = p.add_run(
    "# settings.py\n"
    "DATABASES = {\n"
    "    'default': {\n"
    "        'ENGINE': 'django.db.backends.oracle',\n"
    "        'NAME': 'localhost:1521/XE',\n"
    "        'USER': 'gestion_absences',\n"
    "        'PASSWORD': 'gestion2025',\n"
    "    }\n"
    "}\n\n"
    "# oracle_services.py - Appel d'une fonction Oracle\n"
    "def get_taux_absence(id_etudiant):\n"
    "    with connection.cursor() as cursor:\n"
    "        result = cursor.callfunc(\n"
    "            'PKG_ABSENCES.GET_TAUX_ABSENCE',\n"
    "            float, [id_etudiant])\n"
    "        return result"
)
r.font.name = 'Consolas'
r.font.size = Pt(9)

doc.add_heading("4.3. Interfaces par rôle", level=2)
add_table(
    ["Rôle", "Accès", "Fonctionnalités"],
    [
        ["Administrateur", "Complet", "CRUD étudiants/enseignants/séances, gestion utilisateurs, traitements Oracle, audit, exports, statistiques"],
        ["Enseignant", "Partiel", "Voir ses séances/étudiants, marquer absences, consulter statistiques, exporter"],
        ["Étudiant", "Limité", "Voir ses absences et notifications, soumettre des justificatifs"],
    ]
)

doc.add_page_break()

# ============================================================
# 5. CAPTURES D'ÉCRAN
# ============================================================
doc.add_heading("5. Captures d'écran", level=1)

doc.add_heading("5.1. Page de connexion", level=2)
add_screenshot("page_login", "Figure 1 : Page de connexion de l'application")

doc.add_heading("5.2. Interface Administrateur", level=2)

add_screenshot("admin_dashboard", "Figure 2 : Dashboard administrateur avec KPIs et graphiques (données des vues matérialisées)")
add_screenshot("admin_etudiants", "Figure 3 : Liste des étudiants avec filtres et recherche")
add_screenshot("admin_etudiant_detail", "Figure 4 : Détail d'un étudiant avec statistiques Oracle (fonctions PL/SQL)")
add_screenshot("admin_enseignants", "Figure 5 : Gestion des enseignants")
add_screenshot("admin_seances", "Figure 6 : Liste des séances")
add_screenshot("admin_absences", "Figure 7 : Liste des absences avec filtres avancés")
add_screenshot("admin_marquer_absences", "Figure 8 : Interface de marquage d'absences en masse (appel procédure MARQUER_ABSENCES_BULK)")
add_screenshot("admin_justificatifs", "Figure 9 : Gestion des justificatifs avec documents joints")
add_screenshot("admin_stats_groupes", "Figure 10 : Statistiques par groupe (vue matérialisée MV_STATS_PAR_GROUPE)")
add_screenshot("admin_stats_matieres", "Figure 11 : Statistiques par matière (vue matérialisée MV_STATS_PAR_MATIERE)")
add_screenshot("admin_utilisateurs", "Figure 12 : Gestion des utilisateurs et rôles")
add_screenshot("admin_traitements", "Figure 13 : Traitements Oracle (rafraîchir vues matérialisées, clôturer semestre)")
add_screenshot("admin_audit", "Figure 14 : Journal d'audit alimenté par les triggers Oracle")

doc.add_page_break()
doc.add_heading("5.3. Interface Enseignant", level=2)

add_screenshot("enseignant_dashboard", "Figure 15 : Dashboard enseignant avec ses séances et actions rapides")
add_screenshot("enseignant_absences", "Figure 16 : Liste des absences (filtrées aux séances de l'enseignant)")
add_screenshot("enseignant_etudiants", "Figure 17 : Étudiants des groupes de l'enseignant")
add_screenshot("enseignant_stats", "Figure 18 : Statistiques accessibles à l'enseignant")

doc.add_heading("5.4. Interface Étudiant", level=2)

add_screenshot("etudiant_dashboard", "Figure 19 : Espace étudiant avec ses absences, KPIs et notifications")
add_screenshot("etudiant_absences", "Figure 20 : Mes absences (vue limitée à l'étudiant connecté)")
add_screenshot("etudiant_justificatifs", "Figure 21 : Mes justificatifs soumis")

doc.add_page_break()

# ============================================================
# 6. CONCLUSION
# ============================================================
doc.add_heading("6. Conclusion", level=1)
doc.add_paragraph(
    "Ce projet a permis de mettre en pratique l'ensemble des concepts avancés d'Oracle Database "
    "dans un contexte applicatif réel. La gestion des absences étudiants, bien que simple en apparence, "
    "a nécessité la mise en œuvre de :"
)
doc.add_paragraph("• Un schéma relationnel complet avec 13 tables et toutes les contraintes d'intégrité (PK, FK, UNIQUE, CHECK, NOT NULL)")
doc.add_paragraph("• Une programmation PL/SQL structurée en packages avec 5 fonctions et 6 procédures")
doc.add_paragraph("• 8 triggers pour l'audit automatique et l'application des règles métier")
doc.add_paragraph("• 5 vues et 3 vues matérialisées pour les tableaux de bord et statistiques")
doc.add_paragraph("• Des mécanismes transactionnels (COMMIT, ROLLBACK, SAVEPOINT) pour garantir l'intégrité")
doc.add_paragraph("• Une optimisation via 22 index et des analyses EXPLAIN PLAN")
doc.add_paragraph("• Un volume de données significatif (120 000+ absences) pour valider les performances")
doc.add_paragraph("• Une stratégie de sauvegarde via Data Pump")
doc.add_paragraph(
    "L'intégration avec Django démontre qu'Oracle peut être efficacement utilisé comme moteur de logique "
    "métier (et pas seulement de stockage), les procédures PL/SQL étant appelées directement depuis Python "
    "via cursor.callfunc() et cursor.callproc(). L'application gère 3 rôles utilisateurs (administrateur, "
    "enseignant, étudiant) avec des interfaces adaptées, des exports PDF/Excel, et des graphiques temps réel "
    "alimentés par les vues matérialisées Oracle."
)

# Save
doc.save(OUTPUT)
print(f"Rapport saved to: {OUTPUT}")
